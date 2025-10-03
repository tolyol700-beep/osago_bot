import os
import logging
import io
from datetime import datetime
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import Application, CommandHandler, MessageHandler, ContextTypes, ConversationHandler, filters
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from http.server import HTTPServer, BaseHTTPRequestHandler
from threading import Thread
import time

# ==================== –ü–†–û–°–¢–û–ô –í–ï–ë-–°–ï–†–í–ï–† –î–õ–Ø RENDER ====================
class HealthCheckHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        if self.path == '/':
            self.send_response(200)
            self.send_header('Content-type', 'text/html')
            self.end_headers()
            self.wfile.write(b"""
                <html>
                    <head><title>Insurance Bot</title></head>
                    <body>
                        <h1>ü§ñ –ë–æ—Ç —Å—Ç—Ä–∞—Ö–æ–≤–∞–Ω–∏—è —Ä–∞–±–æ—Ç–∞–µ—Ç!</h1>
                        <p>Insurance Bot is ONLINE and ready to receive applications.</p>
                        <p>üïí –°—Ç–∞—Ç—É—Å: <strong>–ê–∫—Ç–∏–≤–µ–Ω</strong></p>
                        <p>üìÖ –í—Ä–µ–º—è —Å–µ—Ä–≤–µ—Ä–∞: """ + datetime.now().strftime('%Y-%m-%d %H:%M:%S').encode() + b"""</p>
                    </body>
                </html>
            """)
        else:
            self.send_response(404)
            self.end_headers()

def run_health_check():
    port = int(os.environ.get('PORT', 10000))
    server = HTTPServer(('0.0.0.0', port), HealthCheckHandler)
    print(f"‚úÖ –í–µ–±-—Å–µ—Ä–≤–µ—Ä –∑–∞–ø—É—â–µ–Ω –Ω–∞ –ø–æ—Ä—Ç—É {port}")
    server.serve_forever()

# –ó–∞–ø—É—Å–∫–∞–µ–º –≤–µ–±-—Å–µ—Ä–≤–µ—Ä –≤ —Ñ–æ–Ω–µ
health_thread = Thread(target=run_health_check, daemon=True)
health_thread.start()

# ==================== –ó–ê–ì–†–£–ó–ö–ê –ü–ï–†–ï–ú–ï–ù–ù–´–• ====================
load_dotenv()

# ==================== –ù–ê–°–¢–†–û–ô–ö–ê –õ–û–ì–ò–†–û–í–ê–ù–ò–Ø ====================
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

print("üöÄ –ù–∞—á–∏–Ω–∞–µ—Ç—Å—è –∑–∞–ø—É—Å–∫ Telegram –±–æ—Ç–∞...")

# ==================== –°–û–°–¢–û–Ø–ù–ò–Ø –†–ê–ó–ì–û–í–û–†–ê ====================
(
    START, CHOOSE_OWNER_INSURER, INSURER_FIO, INSURER_BIRTHDATE, INSURER_PASSPORT_SERIES_NUMBER,
    INSURER_PASSPORT_ISSUE_DATE, INSURER_PASSPORT_ISSUED_BY, INSURER_PASSPORT_DEPARTMENT_CODE,
    INSURER_REGISTRATION, OWNER_FIO, OWNER_BIRTHDATE, OWNER_PASSPORT_SERIES_NUMBER,
    OWNER_PASSPORT_ISSUE_DATE, OWNER_PASSPORT_ISSUED_BY, OWNER_PASSPORT_DEPARTMENT_CODE,
    INSURER_LICENSE, INSURER_LICENSE_ISSUE_DATE, INSURER_LICENSE_EXPIRY, VEHICLE_BRAND,
    VEHICLE_MODEL, VEHICLE_YEAR, VEHICLE_POWER, VEHICLE_REG_NUMBER, VEHICLE_VIN,
    VEHICLE_DOC_TYPE, VEHICLE_DOC_DETAILS, VEHICLE_DOC_ISSUE_DATE, DRIVERS_CHOICE,
    DRIVER_FIO, DRIVER_LICENSE, DRIVER_LICENSE_ISSUE_DATE, DRIVER_LICENSE_EXPIRY, INSURER_PHONE,
    CONFIRMATION
) = range(34)

user_data = {}

class WordGenerator:
    @staticmethod
    def generate_application_docx(data):
        """–ì–µ–Ω–µ—Ä–∞—Ü–∏—è Word –¥–æ–∫—É–º–µ–Ω—Ç–∞ —Å –∑–∞—è–≤–∫–æ–π"""
        doc = Document()
        
        # –ó–∞–≥–æ–ª–æ–≤–æ–∫
        title = doc.add_heading('–ó–ê–Ø–í–ö–ê –ù–ê –°–¢–†–ê–•–û–í–ê–ù–ò–ï', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # –î–∞—Ç–∞
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"–î–∞—Ç–∞ —Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—è: {datetime.now().strftime('%d.%m.%Y %H:%M')}").bold = True
        doc.add_paragraph()
        
        # –†–∞–∑–¥–µ–ª: –°—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—å
        doc.add_heading('–°–¢–†–ê–•–û–í–ê–¢–ï–õ–¨', level=1)
        
        insurer_info = [
            f"–§–ò–û: {data.get('insurer_fio', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–î–∞—Ç–∞ —Ä–æ–∂–¥–µ–Ω–∏—è: {data.get('insurer_birthdate', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–ü–∞—Å–ø–æ—Ä—Ç: {data.get('insurer_passport_series_number', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞: {data.get('insurer_passport_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–ö–µ–º –≤—ã–¥–∞–Ω: {data.get('insurer_passport_issued_by', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–ö–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è: {data.get('insurer_passport_department_code', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–ü—Ä–æ–ø–∏—Å–∫–∞: {data.get('insurer_registration', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}"
        ]
        
        for info in insurer_info:
            doc.add_paragraph(info)
        
        doc.add_paragraph()
        
        # –†–∞–∑–¥–µ–ª: –°–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫
        doc.add_heading('–°–û–ë–°–¢–í–ï–ù–ù–ò–ö', level=1)
        
        if not data.get('is_same_person', True):
            owner_info = [
                f"–§–ò–û: {data.get('owner_fio', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
                f"–î–∞—Ç–∞ —Ä–æ–∂–¥–µ–Ω–∏—è: {data.get('owner_birthdate', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
                f"–ü–∞—Å–ø–æ—Ä—Ç: {data.get('owner_passport_series_number', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
                f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞: {data.get('owner_passport_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
                f"–ö–µ–º –≤—ã–¥–∞–Ω: {data.get('owner_passport_issued_by', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
                f"–ö–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è: {data.get('owner_passport_department_code', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}"
            ]
            
            for info in owner_info:
                doc.add_paragraph(info)
        else:
            doc.add_paragraph("–°–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫ –∏ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—å - –æ–¥–Ω–æ –ª–∏—Ü–æ")
        
        doc.add_paragraph()
        
        # –í–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–µ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏–µ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è
        doc.add_heading('–í–û–î–ò–¢–ï–õ–¨–°–ö–û–ï –£–î–û–°–¢–û–í–ï–†–ï–ù–ò–ï –°–¢–†–ê–•–û–í–ê–¢–ï–õ–Ø', level=1)
        
        license_info = [
            f"–í/—É: {data.get('insurer_license', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏: {data.get('insurer_license_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–°—Ä–æ–∫ –¥–µ–π—Å—Ç–≤–∏—è: {data.get('insurer_license_expiry', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}"
        ]
        
        for info in license_info:
            doc.add_paragraph(info)
        
        doc.add_paragraph()
        
        # –†–∞–∑–¥–µ–ª: –¢—Ä–∞–Ω—Å–ø–æ—Ä—Ç–Ω–æ–µ —Å—Ä–µ–¥—Å—Ç–≤–æ
        doc.add_heading('–¢–†–ê–ù–°–ü–û–†–¢–ù–û–ï –°–†–ï–î–°–¢–í–û', level=1)
        
        vehicle_info = [
            f"–ú–∞—Ä–∫–∞: {data.get('vehicle_brand', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–ú–æ–¥–µ–ª—å: {data.get('vehicle_model', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–ì–æ–¥ –≤—ã–ø—É—Å–∫–∞: {data.get('vehicle_year', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–ú–æ—â–Ω–æ—Å—Ç—å: {data.get('vehicle_power', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')} –ª.—Å.",
            f"–ì–æ—Å–Ω–æ–º–µ—Ä: {data.get('vehicle_reg_number', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"VIN: {data.get('vehicle_vin', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–î–æ–∫—É–º–µ–Ω—Ç: {data.get('vehicle_doc_type', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')} {data.get('vehicle_doc_details', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}",
            f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {data.get('vehicle_doc_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}"
        ]
        
        for info in vehicle_info:
            doc.add_paragraph(info)
        
        doc.add_paragraph()
        
        # –†–∞–∑–¥–µ–ª: –í–æ–¥–∏—Ç–µ–ª–∏
        doc.add_heading('–í–û–î–ò–¢–ï–õ–ò', level=1)
        
        drivers = data.get('drivers', [])
        if drivers:
            for i, driver in enumerate(drivers, 1):
                driver_paragraph = doc.add_paragraph()
                driver_paragraph.add_run(f'–í–æ–¥–∏—Ç–µ–ª—å {i}: ').bold = True
                driver_paragraph.add_run(f"{driver.get('fio', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}")
                
                doc.add_paragraph(f"   –í/—É: {driver.get('license', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}")
                doc.add_paragraph(f"   –î–∞—Ç–∞ –≤—ã–¥–∞—á–∏: {driver.get('license_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}")
                doc.add_paragraph(f"   –°—Ä–æ–∫ –¥–µ–π—Å—Ç–≤–∏—è: {driver.get('license_expiry', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}")
                doc.add_paragraph()
        else:
            doc.add_paragraph("–í–æ–¥–∏—Ç–µ–ª–∏ –Ω–µ —É–∫–∞–∑–∞–Ω—ã")
        
        # –¢–µ–ª–µ—Ñ–æ–Ω
        doc.add_paragraph()
        phone_paragraph = doc.add_paragraph()
        phone_paragraph.add_run("–¢–µ–ª–µ—Ñ–æ–Ω –¥–ª—è —Å–≤—è–∑–∏: ").bold = True
        phone_paragraph.add_run(f"{data.get('insurer_phone', '–ù–µ —É–∫–∞–∑–∞–Ω')}")
        
        # –ü–æ–¥–ø–∏—Å—å
        doc.add_paragraph()
        doc.add_paragraph("–ó–∞—è–≤–∫–∞ —É—Å–ø–µ—à–Ω–æ –æ—Ñ–æ—Ä–º–ª–µ–Ω–∞!").bold = True
        doc.add_paragraph("–í —Ç–µ—á–µ–Ω–∏–∏ 1 —á–∞—Å–∞ —Å –í–∞–º–∏ —Å–≤—è–∂–µ—Ç—Å—è –º–µ–Ω–µ–¥–∂–µ—Ä, –¥–ª—è –≤–æ–∑–º–æ–∂–Ω–æ–≥–æ —É—Ç–æ—á–Ω–µ–Ω–∏—è –¥–µ—Ç–∞–ª–µ–π –∏ –¥–∞–ª—å–Ω–µ–π—à–µ–≥–æ –æ—Ñ–æ—Ä–º–ª–µ–Ω–∏—è!")
        doc.add_paragraph("–° –£–≤–∞–∂–µ–Ω–∏–µ–º, –ê–û '–ê–ª—å—Ñ–∞—Å—Ç—Ä–∞—Ö–æ–≤–∞–Ω–∏–µ'").bold = True
        
        return doc

def get_navigation_keyboard():
    """–ö–ª–∞–≤–∏–∞—Ç—É—Ä–∞ –¥–ª—è –Ω–∞–≤–∏–≥–∞—Ü–∏–∏"""
    return ReplyKeyboardMarkup([
        ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
    ], resize_keyboard=True)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ù–∞—á–∞–ª–æ —Ä–∞–∑–≥–æ–≤–æ—Ä–∞"""
    user = update.message.from_user
    await update.message.reply_text(
        f"–î–æ–±—Ä–æ –ø–æ–∂–∞–ª–æ–≤–∞—Ç—å, {user.first_name}!\n"
        "–Ø –ø–æ–º–æ–≥—É —Å–æ–±—Ä–∞—Ç—å –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –¥–ª—è —Å—Ç—Ä–∞—Ö–æ–≤–∫–∏.\n\n"
        "–°–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫ –∏ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—å - –æ–¥–Ω–æ –ª–∏—Ü–æ?",
        reply_markup=ReplyKeyboardMarkup([
            ["‚úÖ –û–¥–Ω–æ –ª–∏—Ü–æ", "‚ùå –†–∞–∑–Ω—ã–µ –ª–∏—Ü–∞"]
        ], one_time_keyboard=True, resize_keyboard=True)
    )
    return CHOOSE_OWNER_INSURER

async def choose_owner_insurer(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ –≤—ã–±–æ—Ä–∞ —Ç–∏–ø–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞/—Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text in ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]:
        return await start(update, context)
    
    choice = update.message.text
    user_id = update.message.from_user.id
    
    user_data[user_id] = {
        'is_same_person': choice == "‚úÖ –û–¥–Ω–æ –ª–∏—Ü–æ",
        'drivers': []
    }
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –§–ò–û —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è –ø–æ–ª–Ω–æ—Å—Ç—å—é:",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_FIO

async def insurer_fio(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –§–ò–û —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        return await start(update, context)
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['insurer_fio'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É —Ä–æ–∂–¥–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è (–≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì):\n"
        "–ü—Ä–∏–º–µ—Ä: 15.05.1990",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_BIRTHDATE

async def insurer_birthdate(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞—Ç—ã —Ä–æ–∂–¥–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –§–ò–û —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è –ø–æ–ª–Ω–æ—Å—Ç—å—é:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_FIO
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        user_data[user_id]['insurer_birthdate'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_BIRTHDATE
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:\n"
        "–ü—Ä–∏–º–µ—Ä: 1234 567890",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_PASSPORT_SERIES_NUMBER

async def insurer_passport_series_number(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ —Å–µ—Ä–∏–∏ –∏ –Ω–æ–º–µ—Ä–∞ –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É —Ä–æ–∂–¥–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è (–≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì):",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_BIRTHDATE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['insurer_passport_series_number'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è (–î–î.–ú–ú.–ì–ì–ì–ì):",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_PASSPORT_ISSUE_DATE

async def insurer_passport_issue_date(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞—Ç—ã –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_PASSPORT_SERIES_NUMBER
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        user_data[user_id]['insurer_passport_issue_date'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_PASSPORT_ISSUE_DATE
    
    await update.message.reply_text(
        "–ö–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è?",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_PASSPORT_ISSUED_BY

async def insurer_passport_issued_by(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –æ —Ç–æ–º, –∫–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_PASSPORT_ISSUE_DATE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['insurer_passport_issued_by'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –∫–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_PASSPORT_DEPARTMENT_CODE

async def insurer_passport_department_code(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –∫–æ–¥–∞ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–ö–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è?",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_PASSPORT_ISSUED_BY
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['insurer_passport_department_code'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –ø—Ä–æ–ø–∏—Å–∫—É —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è –≤ —Ñ–æ—Ä–º–∞—Ç–µ:\n"
        "–ì–æ—Ä–æ–¥, –Ω–∞—Å–µ–ª–µ–Ω–Ω—ã–π –ø—É–Ω–∫—Ç, —É–ª–∏—Ü–∞, –¥–æ–º, –∫–æ—Ä–ø—É—Å, –∫–≤–∞—Ä—Ç–∏—Ä–∞\n"
        "–ü—Ä–∏–º–µ—Ä: –ú–æ—Å–∫–≤–∞, –≥. –ú–æ—Å–∫–≤–∞, —É–ª. –õ–µ–Ω–∏–Ω–∞, –¥. 10, –∫. 2, –∫–≤. 25",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_REGISTRATION

async def insurer_registration(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –ø—Ä–æ–ø–∏—Å–∫–∏ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –∫–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_PASSPORT_DEPARTMENT_CODE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['insurer_registration'] = update.message.text
    
    # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω—É–∂–Ω–æ –ª–∏ –∑–∞–ø—Ä–∞—à–∏–≤–∞—Ç—å –¥–∞–Ω–Ω—ã–µ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞
    if user_data[user_id]['is_same_person']:
        # –ï—Å–ª–∏ –æ–¥–Ω–æ –ª–∏—Ü–æ - –ø–µ—Ä–µ—Ö–æ–¥–∏–º –∫ –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–º—É —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—é
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–≥–æ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:\n"
            "–ü—Ä–∏–º–µ—Ä: 1234 567890",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_LICENSE
    else:
        # –ï—Å–ª–∏ —Ä–∞–∑–Ω—ã–µ –ª–∏—Ü–∞ - –∑–∞–ø—Ä–∞—à–∏–≤–∞–µ–º –¥–∞–Ω–Ω—ã–µ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞
        await update.message.reply_text(
            "–¢–µ–ø–µ—Ä—å –≤–≤–µ–¥–µ–º –¥–∞–Ω–Ω—ã–µ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞.\n\n"
            "–í–≤–µ–¥–∏—Ç–µ –§–ò–û —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞ –ø–æ–ª–Ω–æ—Å—Ç—å—é:",
            reply_markup=get_navigation_keyboard()
        )
        return OWNER_FIO

async def owner_fio(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –§–ò–û —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –ø—Ä–æ–ø–∏—Å–∫—É —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_REGISTRATION
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['owner_fio'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É —Ä–æ–∂–¥–µ–Ω–∏—è —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞ (–≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì):\n"
        "–ü—Ä–∏–º–µ—Ä: 15.05.1990",
        reply_markup=get_navigation_keyboard()
    )
    return OWNER_BIRTHDATE

async def owner_birthdate(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞—Ç—ã —Ä–æ–∂–¥–µ–Ω–∏—è —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –§–ò–û —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞ –ø–æ–ª–Ω–æ—Å—Ç—å—é:",
            reply_markup=get_navigation_keyboard()
        )
        return OWNER_FIO
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        user_data[user_id]['owner_birthdate'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return OWNER_BIRTHDATE
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞:\n"
        "–ü—Ä–∏–º–µ—Ä: 1234 567890",
        reply_markup=get_navigation_keyboard()
    )
    return OWNER_PASSPORT_SERIES_NUMBER

async def owner_passport_series_number(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ —Å–µ—Ä–∏–∏ –∏ –Ω–æ–º–µ—Ä–∞ –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É —Ä–æ–∂–¥–µ–Ω–∏—è —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞:",
            reply_markup=get_navigation_keyboard()
        )
        return OWNER_BIRTHDATE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['owner_passport_series_number'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞ (–î–î.–ú–ú.–ì–ì–ì–ì):",
        reply_markup=get_navigation_keyboard()
    )
    return OWNER_PASSPORT_ISSUE_DATE

async def owner_passport_issue_date(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞—Ç—ã –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞:",
            reply_markup=get_navigation_keyboard()
        )
        return OWNER_PASSPORT_SERIES_NUMBER
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        user_data[user_id]['owner_passport_issue_date'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return OWNER_PASSPORT_ISSUE_DATE
    
    await update.message.reply_text(
        "–ö–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞?",
        reply_markup=get_navigation_keyboard()
    )
    return OWNER_PASSPORT_ISSUED_BY

async def owner_passport_issued_by(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –æ —Ç–æ–º, –∫–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞:",
            reply_markup=get_navigation_keyboard()
        )
        return OWNER_PASSPORT_ISSUE_DATE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['owner_passport_issued_by'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –∫–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞:",
        reply_markup=get_navigation_keyboard()
    )
    return OWNER_PASSPORT_DEPARTMENT_CODE

async def owner_passport_department_code(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –∫–æ–¥–∞ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–ö–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞?",
            reply_markup=get_navigation_keyboard()
        )
        return OWNER_PASSPORT_ISSUED_BY
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['owner_passport_department_code'] = update.message.text
    
    # –ü–µ—Ä–µ—Ö–æ–¥–∏–º –∫ –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–º—É —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—é —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–≥–æ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:\n"
        "–ü—Ä–∏–º–µ—Ä: 1234 567890",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_LICENSE

async def insurer_license(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞–Ω–Ω—ã—Ö –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–≥–æ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        if user_data.get(update.message.from_user.id, {}).get('is_same_person', True):
            await update.message.reply_text(
                "–í–≤–µ–¥–∏—Ç–µ –ø—Ä–æ–ø–∏—Å–∫—É —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
                reply_markup=get_navigation_keyboard()
            )
            return INSURER_REGISTRATION
        else:
            await update.message.reply_text(
                "–í–≤–µ–¥–∏—Ç–µ –∫–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –ø–∞—Å–ø–æ—Ä—Ç–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫–∞:",
                reply_markup=get_navigation_keyboard()
            )
            return OWNER_PASSPORT_DEPARTMENT_CODE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['insurer_license'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–≥–æ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è (–î–î.–ú–ú.–ì–ì–ì–ì):",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_LICENSE_ISSUE_DATE

async def insurer_license_issue_date(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞—Ç—ã –≤—ã–¥–∞—á–∏ –ø—Ä–∞–≤ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–≥–æ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_LICENSE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        user_data[user_id]['insurer_license_issue_date'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_LICENSE_ISSUE_DATE
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ —Å—Ä–æ–∫ –æ–∫–æ–Ω—á–∞–Ω–∏—è –¥–µ–π—Å—Ç–≤–∏—è –ø—Ä–∞–≤ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è (–î–î.–ú–ú.–ì–ì–ì–ì):",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_LICENSE_EXPIRY

async def insurer_license_expiry(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ —Å—Ä–æ–∫–∞ –¥–µ–π—Å—Ç–≤–∏—è –ø—Ä–∞–≤ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–≥–æ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_LICENSE_ISSUE_DATE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        user_data[user_id]['insurer_license_expiry'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_LICENSE_EXPIRY
    
    await update.message.reply_text(
        "–¢–µ–ø–µ—Ä—å –≤–≤–µ–¥–µ–º –¥–∞–Ω–Ω—ã–µ —Ç—Ä–∞–Ω—Å–ø–æ—Ä—Ç–Ω–æ–≥–æ —Å—Ä–µ–¥—Å—Ç–≤–∞.\n\n"
        "–í–≤–µ–¥–∏—Ç–µ –º–∞—Ä–∫—É –∞–≤—Ç–æ–º–æ–±–∏–ª—è:",
        reply_markup=get_navigation_keyboard()
    )
    return VEHICLE_BRAND

async def vehicle_brand(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –º–∞—Ä–∫–∏ –∞–≤—Ç–æ–º–æ–±–∏–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Å—Ä–æ–∫ –æ–∫–æ–Ω—á–∞–Ω–∏—è –¥–µ–π—Å—Ç–≤–∏—è –ø—Ä–∞–≤ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_LICENSE_EXPIRY
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['vehicle_brand'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –º–æ–¥–µ–ª—å –∞–≤—Ç–æ–º–æ–±–∏–ª—è:",
        reply_markup=get_navigation_keyboard()
    )
    return VEHICLE_MODEL

async def vehicle_model(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –º–æ–¥–µ–ª–∏ –∞–≤—Ç–æ–º–æ–±–∏–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –º–∞—Ä–∫—É –∞–≤—Ç–æ–º–æ–±–∏–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_BRAND
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['vehicle_model'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –≥–æ–¥ –≤—ã–ø—É—Å–∫–∞ –∞–≤—Ç–æ–º–æ–±–∏–ª—è:",
        reply_markup=get_navigation_keyboard()
    )
    return VEHICLE_YEAR

async def vehicle_year(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –≥–æ–¥–∞ –≤—ã–ø—É—Å–∫–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –º–æ–¥–µ–ª—å –∞–≤—Ç–æ–º–æ–±–∏–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_MODEL
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['vehicle_year'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –º–æ—â–Ω–æ—Å—Ç—å –¥–≤–∏–≥–∞—Ç–µ–ª—è –≤ –ª.—Å.:",
        reply_markup=get_navigation_keyboard()
    )
    return VEHICLE_POWER

async def vehicle_power(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –º–æ—â–Ω–æ—Å—Ç–∏ –¥–≤–∏–≥–∞—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –≥–æ–¥ –≤—ã–ø—É—Å–∫–∞ –∞–≤—Ç–æ–º–æ–±–∏–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_YEAR
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['vehicle_power'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –≥–æ—Å—É–¥–∞—Ä—Å—Ç–≤–µ–Ω–Ω—ã–π –Ω–æ–º–µ—Ä:",
        reply_markup=get_navigation_keyboard()
    )
    return VEHICLE_REG_NUMBER

async def vehicle_reg_number(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –≥–æ—Å –Ω–æ–º–µ—Ä–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –º–æ—â–Ω–æ—Å—Ç—å –¥–≤–∏–≥–∞—Ç–µ–ª—è –≤ –ª.—Å.:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_POWER
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['vehicle_reg_number'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ VIN –Ω–æ–º–µ—Ä:",
        reply_markup=get_navigation_keyboard()
    )
    return VEHICLE_VIN

async def vehicle_vin(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ VIN –Ω–æ–º–µ—Ä–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –≥–æ—Å—É–¥–∞—Ä—Å—Ç–≤–µ–Ω–Ω—ã–π –Ω–æ–º–µ—Ä:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_REG_NUMBER
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['vehicle_vin'] = update.message.text
    
    await update.message.reply_text(
        "–í—ã–±–µ—Ä–∏—Ç–µ —Ç–∏–ø –¥–æ–∫—É–º–µ–Ω—Ç–∞:",
        reply_markup=ReplyKeyboardMarkup([
            ["–°–¢–°", "–ü–¢–°"],
            ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
        ], resize_keyboard=True)
    )
    return VEHICLE_DOC_TYPE

async def vehicle_doc_type(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ —Ç–∏–ø–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ VIN –Ω–æ–º–µ—Ä:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_VIN
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['vehicle_doc_type'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –¥–æ–∫—É–º–µ–Ω—Ç–∞:\n"
        "–ü—Ä–∏–º–µ—Ä: 12–ê–í345678",
        reply_markup=get_navigation_keyboard()
    )
    return VEHICLE_DOC_DETAILS

async def vehicle_doc_details(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ —Å–µ—Ä–∏–∏ –∏ –Ω–æ–º–µ—Ä–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í—ã–±–µ—Ä–∏—Ç–µ —Ç–∏–ø –¥–æ–∫—É–º–µ–Ω—Ç–∞:",
            reply_markup=ReplyKeyboardMarkup([
                ["–°–¢–°", "–ü–¢–°"],
                ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
            ], resize_keyboard=True)
        )
        return VEHICLE_DOC_TYPE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['vehicle_doc_details'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞ (–î–î.–ú–ú.–ì–ì–ì–ì):",
        reply_markup=get_navigation_keyboard()
    )
    return VEHICLE_DOC_ISSUE_DATE

async def vehicle_doc_issue_date(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞—Ç—ã –≤—ã–¥–∞—á–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –¥–æ–∫—É–º–µ–Ω—Ç–∞:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_DOC_DETAILS
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        user_data[user_id]['vehicle_doc_issue_date'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_DOC_ISSUE_DATE
    
    await update.message.reply_text(
        "–¢–µ–ø–µ—Ä—å –¥–æ–±–∞–≤–∏–º –≤–æ–¥–∏—Ç–µ–ª–µ–π.\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ –¥–µ–π—Å—Ç–≤–∏–µ:",
        reply_markup=ReplyKeyboardMarkup([
            ["üìã –°–∫–æ–ø–∏—Ä–æ–≤–∞—Ç—å —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è", "üë§ –î–æ–±–∞–≤–∏—Ç—å –≤–æ–¥–∏—Ç–µ–ª—è"],
            ["‚úÖ –ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ"],
            ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
        ], resize_keyboard=True)
    )
    return DRIVERS_CHOICE

async def drivers_choice(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ –≤—ã–±–æ—Ä–∞ –¥–µ–π—Å—Ç–≤–∏—è —Å –≤–æ–¥–∏—Ç–µ–ª—è–º–∏"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞:",
            reply_markup=get_navigation_keyboard()
        )
        return VEHICLE_DOC_ISSUE_DATE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    choice = update.message.text
    
    if choice == "üìã –°–∫–æ–ø–∏—Ä–æ–≤–∞—Ç—å —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è":
        # –ö–æ–ø–∏—Ä—É–µ–º –¥–∞–Ω–Ω—ã–µ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è
        driver_data = {
            'fio': user_data[user_id]['insurer_fio'],
            'license': user_data[user_id]['insurer_license'],
            'license_issue_date': user_data[user_id]['insurer_license_issue_date'],
            'license_expiry': user_data[user_id]['insurer_license_expiry']
        }
        user_data[user_id]['drivers'].append(driver_data)
        
        await update.message.reply_text(
            "‚úÖ –î–∞–Ω–Ω—ã–µ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è –¥–æ–±–∞–≤–ª–µ–Ω—ã –∫–∞–∫ –≤–æ–¥–∏—Ç–µ–ª—å!\n\n"
            "–í—ã–±–µ—Ä–∏—Ç–µ —Å–ª–µ–¥—É—é—â–µ–µ –¥–µ–π—Å—Ç–≤–∏–µ:",
            reply_markup=ReplyKeyboardMarkup([
                ["üìã –°–∫–æ–ø–∏—Ä–æ–≤–∞—Ç—å —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è", "üë§ –î–æ–±–∞–≤–∏—Ç—å –≤–æ–¥–∏—Ç–µ–ª—è"],
                ["‚úÖ –ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ"],
                ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
            ], resize_keyboard=True)
        )
        return DRIVERS_CHOICE
        
    elif choice == "üë§ –î–æ–±–∞–≤–∏—Ç—å –≤–æ–¥–∏—Ç–µ–ª—è":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –§–ò–û –≤–æ–¥–∏—Ç–µ–ª—è –ø–æ–ª–Ω–æ—Å—Ç—å—é:",
            reply_markup=get_navigation_keyboard()
        )
        return DRIVER_FIO
        
    elif choice == "‚úÖ –ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Ç–µ–ª–µ—Ñ–æ–Ω –¥–ª—è —Å–≤—è–∑–∏:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_PHONE

async def driver_fio(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –§–ò–û –≤–æ–¥–∏—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í—ã–±–µ—Ä–∏—Ç–µ –¥–µ–π—Å—Ç–≤–∏–µ —Å –≤–æ–¥–∏—Ç–µ–ª—è–º–∏:",
            reply_markup=ReplyKeyboardMarkup([
                ["üìã –°–∫–æ–ø–∏—Ä–æ–≤–∞—Ç—å —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è", "üë§ –î–æ–±–∞–≤–∏—Ç—å –≤–æ–¥–∏—Ç–µ–ª—è"],
                ["‚úÖ –ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ"],
                ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
            ], resize_keyboard=True)
        )
        return DRIVERS_CHOICE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    context.user_data['current_driver'] = {'fio': update.message.text}
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–≥–æ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è –≤–æ–¥–∏—Ç–µ–ª—è:\n"
        "–ü—Ä–∏–º–µ—Ä: 1234 567890",
        reply_markup=get_navigation_keyboard()
    )
    return DRIVER_LICENSE

async def driver_license(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –ø—Ä–∞–≤ –≤–æ–¥–∏—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –§–ò–û –≤–æ–¥–∏—Ç–µ–ª—è –ø–æ–ª–Ω–æ—Å—Ç—å—é:",
            reply_markup=get_navigation_keyboard()
        )
        return DRIVER_FIO
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    context.user_data['current_driver']['license'] = update.message.text
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –ø—Ä–∞–≤ (–î–î.–ú–ú.–ì–ì–ì–ì):",
        reply_markup=get_navigation_keyboard()
    )
    return DRIVER_LICENSE_ISSUE_DATE

async def driver_license_issue_date(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞—Ç—ã –≤—ã–¥–∞—á–∏ –ø—Ä–∞–≤ –≤–æ–¥–∏—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Å–µ—Ä–∏—é –∏ –Ω–æ–º–µ—Ä –≤–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–≥–æ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏—è –≤–æ–¥–∏—Ç–µ–ª—è:",
            reply_markup=get_navigation_keyboard()
        )
        return DRIVER_LICENSE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        context.user_data['current_driver']['license_issue_date'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return DRIVER_LICENSE_ISSUE_DATE
    
    await update.message.reply_text(
        "–í–≤–µ–¥–∏—Ç–µ —Å—Ä–æ–∫ –æ–∫–æ–Ω—á–∞–Ω–∏—è –¥–µ–π—Å—Ç–≤–∏—è –ø—Ä–∞–≤ (–î–î.–ú–ú.–ì–ì–ì–ì):",
        reply_markup=get_navigation_keyboard()
    )
    return DRIVER_LICENSE_EXPIRY

async def driver_license_expiry(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ —Å—Ä–æ–∫–∞ –¥–µ–π—Å—Ç–≤–∏—è –ø—Ä–∞–≤ –≤–æ–¥–∏—Ç–µ–ª—è"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏ –ø—Ä–∞–≤:",
            reply_markup=get_navigation_keyboard()
        )
        return DRIVER_LICENSE_ISSUE_DATE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        context.user_data['current_driver']['license_expiry'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:",
            reply_markup=get_navigation_keyboard()
        )
        return DRIVER_LICENSE_EXPIRY
    
    user_data[user_id]['drivers'].append(context.user_data['current_driver'])
    context.user_data.pop('current_driver', None)  # –û—á–∏—â–∞–µ–º –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ
    
    await update.message.reply_text(
        "‚úÖ –í–æ–¥–∏—Ç–µ–ª—å –¥–æ–±–∞–≤–ª–µ–Ω!\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ —Å–ª–µ–¥—É—é—â–µ–µ –¥–µ–π—Å—Ç–≤–∏–µ:",
        reply_markup=ReplyKeyboardMarkup([
            ["üìã –°–∫–æ–ø–∏—Ä–æ–≤–∞—Ç—å —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è", "üë§ –î–æ–±–∞–≤–∏—Ç—å –≤–æ–¥–∏—Ç–µ–ª—è"],
            ["‚úÖ –ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ"],
            ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
        ], resize_keyboard=True)
    )
    return DRIVERS_CHOICE

async def insurer_phone(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–ª—É—á–µ–Ω–∏–µ —Ç–µ–ª–µ—Ñ–æ–Ω–∞ –¥–ª—è —Å–≤—è–∑–∏"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í—ã–±–µ—Ä–∏—Ç–µ –¥–µ–π—Å—Ç–≤–∏–µ —Å –≤–æ–¥–∏—Ç–µ–ª—è–º–∏:",
            reply_markup=ReplyKeyboardMarkup([
                ["üìã –°–∫–æ–ø–∏—Ä–æ–≤–∞—Ç—å —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è", "üë§ –î–æ–±–∞–≤–∏—Ç—å –≤–æ–¥–∏—Ç–µ–ª—è"],
                ["‚úÖ –ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ"],
                ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
            ], resize_keyboard=True)
        )
        return DRIVERS_CHOICE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    if user_id not in user_data:
        await update.message.reply_text("–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –Ω–∞—á–Ω–∏—Ç–µ —Å –∫–æ–º–∞–Ω–¥—ã /start")
        return ConversationHandler.END
        
    user_data[user_id]['insurer_phone'] = update.message.text
    
    # –ü–µ—Ä–µ—Ö–æ–¥ –∫ –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏—é
    await update.message.reply_text(
        "‚úÖ –í—Å–µ –¥–∞–Ω–Ω—ã–µ —Å–æ–±—Ä–∞–Ω—ã!\n\n"
        "–ù–∞–∂–º–∏—Ç–µ –∫–Ω–æ–ø–∫—É –Ω–∏–∂–µ –¥–ª—è –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏—è –∏ –æ—Ç–ø—Ä–∞–≤–∫–∏ –∑–∞—è–≤–∫–∏:",
        reply_markup=ReplyKeyboardMarkup([
            ["‚úÖ –ü–æ–¥—Ç–≤–µ—Ä–¥–∏—Ç—å –∏ –æ—Ç–ø—Ä–∞–≤–∏—Ç—å"],
            ["‚¨ÖÔ∏è –ù–∞–∑–∞–¥", "üè† –í –Ω–∞—á–∞–ª–æ"]
        ], resize_keyboard=True)
    )
    return CONFIRMATION

async def confirmation_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–û–±—Ä–∞–±–æ—Ç–∫–∞ –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏—è –∑–∞—è–≤–∫–∏"""
    # –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞–≤–∏–≥–∞—Ü–∏–æ–Ω–Ω—ã—Ö –∫–Ω–æ–ø–æ–∫
    if update.message.text == "‚¨ÖÔ∏è –ù–∞–∑–∞–¥":
        await update.message.reply_text(
            "–í–≤–µ–¥–∏—Ç–µ —Ç–µ–ª–µ—Ñ–æ–Ω –¥–ª—è —Å–≤—è–∑–∏:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_PHONE
    elif update.message.text == "üè† –í –Ω–∞—á–∞–ª–æ":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    if user_id not in user_data:
        await update.message.reply_text("–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –Ω–∞—á–Ω–∏—Ç–µ —Å –∫–æ–º–∞–Ω–¥—ã /start")
        return ConversationHandler.END
        
    # –í—ã–∑—ã–≤–∞–µ–º —Ñ—É–Ω–∫—Ü–∏—é –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏—è –∏ –æ—Ç–ø—Ä–∞–≤–∫–∏
    return await send_confirmation(update, context)

async def send_confirmation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–ü–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏–µ –∏ –æ—Ç–ø—Ä–∞–≤–∫–∞ –¥–∞–Ω–Ω—ã—Ö"""
    user_id = update.message.from_user.id
    if user_id not in user_data:
        await update.message.reply_text("–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –Ω–∞—á–Ω–∏—Ç–µ —Å –∫–æ–º–∞–Ω–¥—ã /start")
        return ConversationHandler.END
        
    data = user_data[user_id]
    
    try:
        # –§–æ—Ä–º–∏—Ä—É–µ–º –¥–µ—Ç–∞–ª—å–Ω–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ –¥–ª—è Telegram
        manager_message = "üöó –°–†–û–ß–ù–ê–Ø –ó–ê–Ø–í–ö–ê –ù–ê –°–¢–†–ê–•–û–í–ê–ù–ò–ï\n\n"
        
        manager_message += "üë§ –°–¢–†–ê–•–û–í–ê–¢–ï–õ–¨:\n"
        manager_message += f"–§–ò–û: {data.get('insurer_fio', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–î–∞—Ç–∞ —Ä–æ–∂–¥–µ–Ω–∏—è: {data.get('insurer_birthdate', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–ü–∞—Å–ø–æ—Ä—Ç: {data.get('insurer_passport_series_number', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏: {data.get('insurer_passport_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–ö–µ–º –≤—ã–¥–∞–Ω: {data.get('insurer_passport_issued_by', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–ö–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è: {data.get('insurer_passport_department_code', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–ü—Ä–æ–ø–∏—Å–∫–∞: {data.get('insurer_registration', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n\n"
        
        # –î–û–ë–ê–í–õ–ï–ù–û: –í–æ–¥–∏—Ç–µ–ª—å—Å–∫–æ–µ —É–¥–æ—Å—Ç–æ–≤–µ—Ä–µ–Ω–∏–µ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—è
        manager_message += "üöó –í–û–î–ò–¢–ï–õ–¨–°–ö–û–ï –£–î–û–°–¢–û–í–ï–†–ï–ù–ò–ï –°–¢–†–ê–•–û–í–ê–¢–ï–õ–Ø:\n"
        manager_message += f"–ù–æ–º–µ—Ä: {data.get('insurer_license', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏: {data.get('insurer_license_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–°—Ä–æ–∫ –¥–µ–π—Å—Ç–≤–∏—è: {data.get('insurer_license_expiry', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n\n"
        
        if not data.get('is_same_person', True):
            manager_message += "üë§ –°–û–ë–°–¢–í–ï–ù–ù–ò–ö:\n"
            manager_message += f"–§–ò–û: {data.get('owner_fio', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
            manager_message += f"–î–∞—Ç–∞ —Ä–æ–∂–¥–µ–Ω–∏—è: {data.get('owner_birthdate', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
            manager_message += f"–ü–∞—Å–ø–æ—Ä—Ç: {data.get('owner_passport_series_number', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
            manager_message += f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏: {data.get('owner_passport_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
            manager_message += f"–ö–µ–º –≤—ã–¥–∞–Ω: {data.get('owner_passport_issued_by', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
            manager_message += f"–ö–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è: {data.get('owner_passport_department_code', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n\n"
        else:
            manager_message += "üë§ –°–û–ë–°–¢–í–ï–ù–ù–ò–ö:\n"
            manager_message += "–°–æ–±—Å—Ç–≤–µ–Ω–Ω–∏–∫ –∏ —Å—Ç—Ä–∞—Ö–æ–≤–∞—Ç–µ–ª—å - –æ–¥–Ω–æ –ª–∏—Ü–æ\n\n"
        
        manager_message += "üöó –¢–†–ê–ù–°–ü–û–†–¢–ù–û–ï –°–†–ï–î–°–¢–í–û:\n"
        manager_message += f"–ú–∞—Ä–∫–∞: {data.get('vehicle_brand', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–ú–æ–¥–µ–ª—å: {data.get('vehicle_model', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–ì–æ–¥ –≤—ã–ø—É—Å–∫–∞: {data.get('vehicle_year', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–ú–æ—â–Ω–æ—Å—Ç—å: {data.get('vehicle_power', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')} –ª.—Å.\n"
        manager_message += f"–ì–æ—Å–Ω–æ–º–µ—Ä: {data.get('vehicle_reg_number', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"VIN: {data.get('vehicle_vin', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–î–æ–∫—É–º–µ–Ω—Ç: {data.get('vehicle_doc_type', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')} {data.get('vehicle_doc_details', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
        manager_message += f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏: {data.get('vehicle_doc_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n\n"
        
        manager_message += "üë• –í–û–î–ò–¢–ï–õ–ò:\n"
        drivers = data.get('drivers', [])
        if drivers:
            for i, driver in enumerate(drivers, 1):
                manager_message += f"{i}. {driver.get('fio', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
                manager_message += f"   –í/—É: {driver.get('license', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
                manager_message += f"   –î–∞—Ç–∞ –≤—ã–¥–∞—á–∏: {driver.get('license_issue_date', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n"
                manager_message += f"   –°—Ä–æ–∫ –¥–µ–π—Å—Ç–≤–∏—è: {driver.get('license_expiry', '–ù–µ —É–∫–∞–∑–∞–Ω–æ')}\n\n"
        else:
            manager_message += "–í–æ–¥–∏—Ç–µ–ª–∏ –Ω–µ —É–∫–∞–∑–∞–Ω—ã\n\n"
        
        manager_message += f"üìû –¢–µ–ª–µ—Ñ–æ–Ω: {data.get('insurer_phone', '–ù–µ —É–∫–∞–∑–∞–Ω')}\n"
        manager_message += f"üìÖ –î–∞—Ç–∞ –∑–∞—è–≤–∫–∏: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        
        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –¥–µ—Ç–∞–ª—å–Ω–æ–µ —É–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ –º–µ–Ω–µ–¥–∂–µ—Ä—É –≤ Telegram
        MANAGER_CHAT_ID = os.getenv('MANAGER_CHAT_ID')
        if MANAGER_CHAT_ID:
            try:
                # –†–∞–∑–±–∏–≤–∞–µ–º –¥–ª–∏–Ω–Ω–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ –Ω–∞ —á–∞—Å—Ç–∏ –µ—Å–ª–∏ –Ω—É–∂–Ω–æ
                if len(manager_message) > 4096:
                    parts = [manager_message[i:i+4096] for i in range(0, len(manager_message), 4096)]
                    for part in parts:
                        await context.bot.send_message(chat_id=int(MANAGER_CHAT_ID), text=part)
                else:
                    await context.bot.send_message(chat_id=int(MANAGER_CHAT_ID), text=manager_message)
                
                print(f"‚úÖ –¢–µ–∫—Å—Ç–æ–≤–æ–µ —É–≤–µ–¥–æ–º–ª–µ–Ω–∏–µ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω–æ –º–µ–Ω–µ–¥–∂–µ—Ä—É {MANAGER_CHAT_ID}")
            except Exception as e:
                print(f"‚ùå –û—à–∏–±–∫–∞ –æ—Ç–ø—Ä–∞–≤–∫–∏ –≤ Telegram: {e}")
        
        # –°–æ–∑–¥–∞–µ–º Word –¥–æ–∫—É–º–µ–Ω—Ç
        doc = WordGenerator.generate_application_docx(data)
        
        # –°–æ—Ö—Ä–∞–Ω—è–µ–º Word –≤ –±–∞–π—Ç—ã
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        file_stream.name = f"–ó–∞—è–≤–∫–∞_{data.get('insurer_fio', '–ö–ª–∏–µ–Ω—Ç')}_{datetime.now().strftime('%d%m%Y_%H%M')}.docx"
        
        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º Word –¥–æ–∫—É–º–µ–Ω—Ç –º–µ–Ω–µ–¥–∂–µ—Ä—É
        if MANAGER_CHAT_ID:
            try:
                await context.bot.send_document(
                    chat_id=int(MANAGER_CHAT_ID),
                    document=file_stream,
                    caption=f"üìÑ –ó–∞—è–≤–∫–∞ –æ—Ç {data.get('insurer_fio', '–ö–ª–∏–µ–Ω—Ç')}"
                )
                print(f"‚úÖ Word –¥–æ–∫—É–º–µ–Ω—Ç –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω –º–µ–Ω–µ–¥–∂–µ—Ä—É {MANAGER_CHAT_ID}")
            except Exception as e:
                print(f"‚ùå –û—à–∏–±–∫–∞ –æ—Ç–ø—Ä–∞–≤–∫–∏ Word –º–µ–Ω–µ–¥–∂–µ—Ä—É: {e}")
        
        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏–µ –∫–ª–∏–µ–Ω—Ç—É
        await update.message.reply_text(
            "‚úÖ –ó–∞—è–≤–∫–∞ —É—Å–ø–µ—à–Ω–æ –æ—Ñ–æ—Ä–º–ª–µ–Ω–∞!\n\n"
            "–í —Ç–µ—á–µ–Ω–∏–∏ 1 —á–∞—Å–∞ —Å –í–∞–º–∏ —Å–≤—è–∂–µ—Ç—Å—è –º–µ–Ω–µ–¥–∂–µ—Ä, –¥–ª—è –≤–æ–∑–º–æ–∂–Ω–æ–≥–æ —É—Ç–æ—á–Ω–µ–Ω–∏—è –¥–µ—Ç–∞–ª–µ–π –∏ –¥–∞–ª—å–Ω–µ–π—à–µ–≥–æ –æ—Ñ–æ—Ä–º–ª–µ–Ω–∏—è!\n\n"
            "–° –£–≤–∞–∂–µ–Ω–∏–µ–º, –ê–û '–ê–ª—å—Ñ–∞—Å—Ç—Ä–∞—Ö–æ–≤–∞–Ω–∏–µ'",
            reply_markup=ReplyKeyboardRemove()
        )
        
        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Ç–µ–∫—Å—Ç–æ–≤—É—é –∫–æ–ø–∏—é –∫–ª–∏–µ–Ω—Ç—É (–° –î–û–ë–ê–í–õ–ï–ù–ù–´–ú–ò –î–ê–ù–ù–´–ú–ò –í–û–î–ò–¢–ï–õ–¨–°–ö–û–ì–û –£–î–û–°–¢–û–í–ï–†–ï–ù–ò–Ø)
        client_message = "üìã –í–∞—à–∞ –∑–∞—è–≤–∫–∞:\n\n" + manager_message
        if len(client_message) > 4096:
            parts = [client_message[i:i+4096] for i in range(0, len(client_message), 4096)]
            for part in parts:
                await update.message.reply_text(part)
        else:
            await update.message.reply_text(client_message)
        
        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º Word –¥–æ–∫—É–º–µ–Ω—Ç –∫–ª–∏–µ–Ω—Ç—É
        file_stream.seek(0)  # –°–±—Ä–∞—Å—ã–≤–∞–µ–º –ø–æ–∑–∏—Ü–∏—é –¥–ª—è –ø–æ–≤—Ç–æ—Ä–Ω–æ–≥–æ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è
        await update.message.reply_document(
            document=file_stream,
            caption="üìÑ –í–∞—à–∞ –∑–∞—è–≤–∫–∞ –Ω–∞ —Å—Ç—Ä–∞—Ö–æ–≤–∞–Ω–∏–µ"
        )
        
    except Exception as e:
        print(f"‚ùå –ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞: {e}")
        await update.message.reply_text(
            "–ü—Ä–æ–∏–∑–æ—à–ª–∞ –Ω–µ–ø—Ä–µ–¥–≤–∏–¥–µ–Ω–Ω–∞—è –æ—à–∏–±–∫–∞. "
            "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –ø–æ–ø—Ä–æ–±—É–π—Ç–µ –ø–æ–∑–∂–µ.",
            reply_markup=ReplyKeyboardRemove()
        )
    
    # –û—á–∏—â–∞–µ–º –¥–∞–Ω–Ω—ã–µ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
    if user_id in user_data:
        del user_data[user_id]
    
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """–û—Ç–º–µ–Ω–∞ —Ä–∞–∑–≥–æ–≤–æ—Ä–∞"""
    user_id = update.message.from_user.id
    if user_id in user_data:
        del user_data[user_id]
    
    await update.message.reply_text(
        "–ó–∞—è–≤–∫–∞ –æ—Ç–º–µ–Ω–µ–Ω–∞.",
        reply_markup=ReplyKeyboardRemove()
    )
    return ConversationHandler.END

def main():
    """–ó–∞–ø—É—Å–∫ –±–æ—Ç–∞"""
    TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
    
    if not TOKEN:
        logging.error("‚ùå –û—à–∏–±–∫–∞: –Ω–µ –∑–∞–¥–∞–Ω TELEGRAM_BOT_TOKEN")
        return
    
    try:
        application = Application.builder().token(TOKEN).build()
        
        conv_handler = ConversationHandler(
            entry_points=[CommandHandler('start', start)],
            states={
                CHOOSE_OWNER_INSURER: [MessageHandler(filters.TEXT & ~filters.COMMAND, choose_owner_insurer)],
                INSURER_FIO: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_fio)],
                INSURER_BIRTHDATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_birthdate)],
                INSURER_PASSPORT_SERIES_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_passport_series_number)],
                INSURER_PASSPORT_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_passport_issue_date)],
                INSURER_PASSPORT_ISSUED_BY: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_passport_issued_by)],
                INSURER_PASSPORT_DEPARTMENT_CODE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_passport_department_code)],
                INSURER_REGISTRATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_registration)],
                OWNER_FIO: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_fio)],
                OWNER_BIRTHDATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_birthdate)],
                OWNER_PASSPORT_SERIES_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_passport_series_number)],
                OWNER_PASSPORT_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_passport_issue_date)],
                OWNER_PASSPORT_ISSUED_BY: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_passport_issued_by)],
                OWNER_PASSPORT_DEPARTMENT_CODE: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_passport_department_code)],
                INSURER_LICENSE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_license)],
                INSURER_LICENSE_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_license_issue_date)],
                INSURER_LICENSE_EXPIRY: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_license_expiry)],
                VEHICLE_BRAND: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_brand)],
                VEHICLE_MODEL: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_model)],
                VEHICLE_YEAR: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_year)],
                VEHICLE_POWER: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_power)],
                VEHICLE_REG_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_reg_number)],
                VEHICLE_VIN: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_vin)],
                VEHICLE_DOC_TYPE: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_doc_type)],
                VEHICLE_DOC_DETAILS: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_doc_details)],
                VEHICLE_DOC_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_doc_issue_date)],
                DRIVERS_CHOICE: [MessageHandler(filters.TEXT & ~filters.COMMAND, drivers_choice)],
                DRIVER_FIO: [MessageHandler(filters.TEXT & ~filters.COMMAND, driver_fio)],
                DRIVER_LICENSE: [MessageHandler(filters.TEXT & ~filters.COMMAND, driver_license)],
                DRIVER_LICENSE_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, driver_license_issue_date)],
                DRIVER_LICENSE_EXPIRY: [MessageHandler(filters.TEXT & ~filters.COMMAND, driver_license_expiry)],
                INSURER_PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_phone)],
                CONFIRMATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, confirmation_handler)],
            },
            fallbacks=[CommandHandler('cancel', cancel)]
        )
        
        application.add_handler(conv_handler)
        
        logging.info("ü§ñ –ë–æ—Ç –∑–∞–ø—É—Å–∫–∞–µ—Ç—Å—è...")
        print("=== –ë–û–¢ –ó–ê–ü–£–©–ï–ù –ù–ê RENDER ===")
        
        application.run_polling(
            drop_pending_updates=True,
            allowed_updates=Update.ALL_TYPES,
            close_loop=False
        )
        
    except Exception as e:
        logging.error(f"‚ùå –ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞: {e}")
        print("–ë–æ—Ç –æ—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –∏–∑-–∑–∞ –æ—à–∏–±–∫–∏:", e)
        # –ü–µ—Ä–µ–∑–∞–ø—É—Å–∫ —á–µ—Ä–µ–∑ 10 —Å–µ–∫—É–Ω–¥
        time.sleep(10)
        main()

if __name__ == '__main__':
    main()
